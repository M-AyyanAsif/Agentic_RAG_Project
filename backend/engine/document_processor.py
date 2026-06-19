"""
Document parsing and Pinecone Indexing (RAG-ready version).
- Production Optimized: Tailored for gemini-embedding-2 with safe string formatting.
- Memory Safe: Uses batch upserting to prevent RAM exhaustion.
- Quota Safe: Implements proactive pacing and aggressive backoff for API limits.
"""

from __future__ import annotations
import io
import logging
import re
import time
from dataclasses import dataclass
from typing import List
import docx
from pypdf import PdfReader
from google import genai
from google.genai import types
from pinecone import Pinecone

from backend.core.config import get_settings

logger = logging.getLogger(__name__)

@dataclass(slots=True)
class ParsedDocument:
    filename: str
    content: str
    content_type: str
    chunks: List[str]

class DocumentProcessor:
    SUPPORTED_TYPES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    def __init__(self, chunk_size: int = 2000, chunk_overlap: int = 300):
        self.settings = get_settings()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize Modern Google GenAI Client
        self.client = genai.Client(api_key=self.settings.google_api_key)
        
        # Initialize Pinecone Cloud Client
        self.pc = Pinecone(api_key=self.settings.pinecone_api_key)
        self.index = self.pc.Index(self.settings.pinecone_index_name)

    def parse(self, filename: str, content_type: str, file_bytes: bytes) -> ParsedDocument:
        if content_type not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {content_type}")

        if content_type == "application/pdf":
            text = self._parse_pdf(file_bytes)
        else:
            text = self._parse_docx(file_bytes)

        cleaned_text = self._clean_text(text)
        
        if not cleaned_text:
            raise ValueError(
                f"Text extraction failed for '{filename}'. Document is empty, encrypted, or holds unreadable scanned imagery."
            )

        chunks = self._chunk_text(cleaned_text)

        return ParsedDocument(
            filename=filename,
            content=cleaned_text,
            content_type=content_type,
            chunks=chunks,
        )

    def index_to_pinecone(self, session_id: str, parsed_doc: ParsedDocument) -> None:
        """
        Computes embeddings via Gemini API and streams updates to Pinecone Cloud Indexes.
        Includes aggressive backoff and pacing to prevent 429 RESOURCE_EXHAUSTED errors.
        """
        try:
            vectors_to_upsert = []
            batch_size = 100  
            total_chunks = len(parsed_doc.chunks)
            
            logger.info(f"Generating embeddings for {total_chunks} chunks from '{parsed_doc.filename}' via gemini-embedding-2...")

            for i, chunk in enumerate(parsed_doc.chunks):
                # UPGRADED: 5 Retries, starting at a massive 15-second delay to wait out minute limits
                retries = 5
                delay = 15.0 
                res = None
                
                formatted_content = f"title: {parsed_doc.filename} | text: {chunk}"
                
                for attempt in range(retries):
                    try:
                        res = self.client.models.embed_content(
                            model="gemini-embedding-2",
                            contents=formatted_content,
                            config=types.EmbedContentConfig(
                                task_type="RETRIEVAL_DOCUMENT", 
                                output_dimensionality=1024,
                                title=parsed_doc.filename
                            )
                        )
                        break  
                    except Exception as exc:
                        err_msg = str(exc)
                        if "429" in err_msg or "Resource exhausted" in err_msg:
                            if attempt == retries - 1:
                                logger.error(f"Failed completely after {retries} retries. Quota is completely exhausted.")
                                raise exc
                            
                            logger.warning(f"Rate limit hit on chunk {i+1}/{total_chunks}. Backing off for {delay} seconds to let quota reset...")
                            time.sleep(delay)
                            delay = min(delay * 1.5, 60.0) # Scale up delay, but cap at 60 seconds 
                        else:
                            logger.error(f"Skipping chunk index {i} due to unhandled API exception: {exc}")
                            break  

                if res and hasattr(res, 'embeddings') and res.embeddings:
                    embedding = res.embeddings[0].values
                    
                    clean_filename = re.sub(r'[^a-zA-Z0-9_\-]', '', parsed_doc.filename.replace(' ', '_'))
                    vector_id = f"{session_id}_{clean_filename}_{i}"
                    
                    vectors_to_upsert.append({
                        "id": vector_id,
                        "values": embedding,
                        "metadata": {
                            "text": chunk,
                            "source": parsed_doc.filename, 
                            "filename": parsed_doc.filename,
                            "session_id": session_id
                        }
                    })
                
                if len(vectors_to_upsert) >= batch_size:
                    self.index.upsert(vectors=vectors_to_upsert, namespace=session_id)
                    logger.info(f"Successfully flushed batch of {len(vectors_to_upsert)} vectors to Pinecone ({i+1}/{total_chunks}).")
                    vectors_to_upsert.clear()  
                
                # UPGRADED: Proactive Pacing. Pause for 3 seconds between every chunk to artificially limit RPM
                time.sleep(3.0)

            if vectors_to_upsert:
                self.index.upsert(vectors=vectors_to_upsert, namespace=session_id)
                logger.info(f"Successfully flushed remaining {len(vectors_to_upsert)} trailing vectors to cloud index.")
                vectors_to_upsert.clear()
                
            logger.info(f"RAG Compilation Success: All processed blocks compiled inside namespace: {session_id}")
            
        except Exception as e:
            logger.error(f"Pinecone deployment runtime exception: {e}")
            raise e

    def _parse_pdf(self, file_bytes: bytes) -> str:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            text_blocks = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_blocks.append(text)
            return "\n\n".join(text_blocks)
        except Exception as e:
            logger.error("PDF engine exception: %s", str(e))
            return ""

    def _parse_docx(self, file_bytes: bytes) -> str:
        try:
            document = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error("DOCX engine exception: %s", str(e))
            return ""

    def _clean_text(self, text: str) -> str:
        text = re.sub(r"[ \t]+", " ", text) 
        text = re.sub(r"\n\s*\n", "\n\n", text)
        return text.strip()

    def _chunk_text(self, text: str) -> List[str]:
        if not text:
            return []
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            if end < len(text):
                last_space = text.rfind(' ', start, end)
                if last_space != -1 and last_space > start: 
                    end = last_space
            
            chunk_slice = text[start:end].strip()
            if chunk_slice:
                chunks.append(chunk_slice)
            
            next_start = end - self.chunk_overlap
            if next_start <= start:
                start = end
            else:
                start = next_start
                
        return [c for c in chunks if len(c) > 10]